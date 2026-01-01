from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Dict, Any, List
import io
import json


class DOCXGenerator:
    """Generate professional DOCX reports following Indian engineering college standards (GTU/VTU/AICTE)"""
    
    def _ensure_dict(self, data: Any, default: Any = None) -> Dict:
        """Ensure data is a dictionary, parsing JSON if needed"""
        if data is None:
            return default if default is not None else {}
        if isinstance(data, str):
            try:
                parsed = json.loads(data)
                return parsed if isinstance(parsed, dict) else (default if default is not None else {})
            except json.JSONDecodeError:
                return default if default is not None else {}
        if isinstance(data, dict):
            return data
        return default if default is not None else {}
    
    def _ensure_list(self, data: Any) -> List:
        """Ensure data is a list"""
        if data is None:
            return []
        if isinstance(data, list):
            return data
        if isinstance(data, str):
            return [data]
        return []
    
    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """Add a properly formatted heading"""
        heading = doc.add_heading(text, level)
        return heading
    
    def _add_table_from_data(self, doc: Document, headers: List[str], rows: List[List[str]]):
        """Create a formatted table"""
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Data rows
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(row_cells):
                    row_cells[col_idx].text = str(cell_data)
        
        return table
    
    def generate_report(self, project_data: Dict[str, Any]) -> bytes:
        """
        Generate professional DOCX report following Indian engineering standards
        
        Args:
            project_data: Project JSON from LLM
            
        Returns:
            DOCX file as bytes
        """
        project_data = self._ensure_dict(project_data)
        doc = Document()
        
        # ============ TITLE PAGE ============
        title = doc.add_heading(project_data.get('title', 'Project Report'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        subtitle = doc.add_paragraph("A Project Report")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph("Submitted in partial fulfillment of the requirements for the degree of")
        doc.add_paragraph("Bachelor of Engineering / Diploma in Engineering")
        doc.add_paragraph()
        
        # Add difficulty and timeline info
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"Difficulty: {project_data.get('difficulty', 'Intermediate')}").bold = True
        info_para.add_run(f" | Timeline: {project_data.get('timeline_days', 45)} days")
        info_para.add_run(f" | Estimated LOC: {project_data.get('estimated_loc', 2000)}")
        
        # ============ ABSTRACT ============
        doc.add_page_break()
        doc.add_heading('ABSTRACT', 1)
        
        abstract = project_data.get('abstract', '')
        if isinstance(abstract, str):
            doc.add_paragraph(abstract)
        
        # Keywords
        keywords = project_data.get('keywords', [])
        if keywords:
            doc.add_paragraph()
            kw_para = doc.add_paragraph()
            kw_para.add_run("Keywords: ").bold = True
            kw_para.add_run(', '.join(keywords))
        
        # ============ TABLE OF CONTENTS (Placeholder) ============
        doc.add_page_break()
        doc.add_heading('TABLE OF CONTENTS', 1)
        
        toc_items = [
            ("Abstract", "i"),
            ("List of Figures", "ii"),
            ("List of Tables", "iii"),
            ("Chapter 1: Introduction", "1"),
            ("Chapter 2: Literature Survey", "5"),
            ("Chapter 3: System Study", "10"),
            ("Chapter 4: System Design", "15"),
            ("Chapter 5: Implementation", "22"),
            ("Chapter 6: Testing", "28"),
            ("Chapter 7: Results & Screenshots", "32"),
            ("Chapter 8: Conclusion & Future Scope", "36"),
            ("References", "38"),
            ("Appendix A: Source Code", "40"),
        ]
        
        for item, page in toc_items:
            para = doc.add_paragraph()
            para.add_run(item)
            para.add_run('\t' * 8)
            para.add_run(page)
        
        # ============ CHAPTER 1: INTRODUCTION ============
        doc.add_page_break()
        doc.add_heading('CHAPTER 1: INTRODUCTION', 1)
        
        introduction = self._ensure_dict(project_data.get('introduction', {}))
        if isinstance(project_data.get('introduction'), str):
            doc.add_paragraph(project_data.get('introduction', ''))
        else:
            if introduction.get('background'):
                doc.add_heading('1.1 Background', 2)
                doc.add_paragraph(introduction.get('background', ''))
            
            if introduction.get('problem_statement'):
                doc.add_heading('1.2 Problem Statement', 2)
                doc.add_paragraph(introduction.get('problem_statement', ''))
            
            if introduction.get('motivation'):
                doc.add_heading('1.3 Motivation', 2)
                doc.add_paragraph(introduction.get('motivation', ''))
            
            if introduction.get('project_overview'):
                doc.add_heading('1.4 Project Overview', 2)
                doc.add_paragraph(introduction.get('project_overview', ''))
        
        # Objectives
        objectives = self._ensure_list(project_data.get('objectives', []))
        if objectives:
            doc.add_heading('1.5 Objectives', 2)
            for obj in objectives:
                doc.add_paragraph(obj, style='List Bullet')
        
        # Scope
        scope = project_data.get('scope', {})
        if scope:
            doc.add_heading('1.6 Scope', 2)
            if isinstance(scope, str):
                doc.add_paragraph(scope)
            elif isinstance(scope, dict):
                if scope.get('in_scope'):
                    doc.add_heading('In Scope:', 3)
                    for item in self._ensure_list(scope.get('in_scope', [])):
                        doc.add_paragraph(item, style='List Bullet')
                if scope.get('out_of_scope'):
                    doc.add_heading('Out of Scope:', 3)
                    for item in self._ensure_list(scope.get('out_of_scope', [])):
                        doc.add_paragraph(item, style='List Bullet')
                if scope.get('target_users'):
                    doc.add_paragraph(f"Target Users: {scope.get('target_users')}")
        
        # ============ CHAPTER 2: LITERATURE SURVEY ============
        doc.add_page_break()
        doc.add_heading('CHAPTER 2: LITERATURE SURVEY', 1)
        
        literature = self._ensure_list(project_data.get('literature_survey', []))
        if literature:
            for idx, paper in enumerate(literature, 1):
                if isinstance(paper, dict):
                    doc.add_heading(f"2.{idx} {paper.get('paper_title', 'Research Paper')}", 2)
                    doc.add_paragraph(f"Authors: {paper.get('authors', 'N/A')}")
                    doc.add_paragraph(f"Year: {paper.get('year', 'N/A')}")
                    doc.add_paragraph(f"Publication: {paper.get('publication', 'N/A')}")
                    doc.add_paragraph()
                    doc.add_paragraph(f"Summary: {paper.get('summary', '')}")
                    doc.add_paragraph(f"Relevance: {paper.get('relevance', '')}")
                    if paper.get('gap_identified'):
                        doc.add_paragraph(f"Gap Identified: {paper.get('gap_identified', '')}")
                elif isinstance(paper, str):
                    doc.add_paragraph(paper, style='List Number')
        else:
            doc.add_paragraph("Literature survey to be completed based on domain research.")
        
        # ============ CHAPTER 3: SYSTEM STUDY ============
        doc.add_page_break()
        doc.add_heading('CHAPTER 3: SYSTEM STUDY', 1)
        
        system_study = self._ensure_dict(project_data.get('system_study', {}))
        
        existing = self._ensure_dict(system_study.get('existing_system', {}))
        if existing:
            doc.add_heading('3.1 Existing System', 2)
            doc.add_paragraph(existing.get('description', ''))
            
            if existing.get('limitations'):
                doc.add_heading('Limitations of Existing System:', 3)
                for lim in self._ensure_list(existing.get('limitations', [])):
                    doc.add_paragraph(lim, style='List Bullet')
        
        proposed = self._ensure_dict(system_study.get('proposed_system', {}))
        if proposed:
            doc.add_heading('3.2 Proposed System', 2)
            doc.add_paragraph(proposed.get('description', ''))
            
            if proposed.get('advantages'):
                doc.add_heading('Advantages:', 3)
                for adv in self._ensure_list(proposed.get('advantages', [])):
                    doc.add_paragraph(adv, style='List Bullet')
        
        # ============ CHAPTER 4: SYSTEM DESIGN ============
        doc.add_page_break()
        doc.add_heading('CHAPTER 4: SYSTEM DESIGN', 1)
        
        system_design = self._ensure_dict(project_data.get('system_design', {}))
        
        # Architecture
        arch = self._ensure_dict(system_design.get('architecture', {}))
        if arch:
            doc.add_heading('4.1 System Architecture', 2)
            doc.add_paragraph(f"Architecture Type: {arch.get('type', 'N/A')}")
            doc.add_paragraph(arch.get('description', ''))
            if arch.get('diagram_description'):
                doc.add_paragraph(f"[Architecture Diagram: {arch.get('diagram_description')}]")
        
        # Use Case Diagram
        use_case = self._ensure_dict(system_design.get('use_case', {}))
        if use_case:
            doc.add_heading('4.2 Use Case Diagram', 2)
            if use_case.get('actors'):
                doc.add_paragraph(f"Actors: {', '.join(self._ensure_list(use_case.get('actors', [])))}")
            if use_case.get('use_cases'):
                doc.add_heading('Use Cases:', 3)
                for uc in self._ensure_list(use_case.get('use_cases', [])):
                    if isinstance(uc, dict):
                        doc.add_paragraph(f"• {uc.get('name', '')}: {uc.get('description', '')}")
        
        # Class Diagram
        class_diag = self._ensure_dict(system_design.get('class_diagram', {}))
        if class_diag and class_diag.get('classes'):
            doc.add_heading('4.3 Class Diagram', 2)
            for cls in self._ensure_list(class_diag.get('classes', [])):
                if isinstance(cls, dict):
                    doc.add_paragraph(f"Class: {cls.get('name', 'ClassName')}")
                    if cls.get('attributes'):
                        doc.add_paragraph(f"  Attributes: {', '.join(self._ensure_list(cls.get('attributes', [])))}")
                    if cls.get('methods'):
                        doc.add_paragraph(f"  Methods: {', '.join(self._ensure_list(cls.get('methods', [])))}")
        
        # ER Diagram
        er_diag = self._ensure_dict(system_design.get('er_diagram', {}))
        if er_diag:
            doc.add_heading('4.4 ER Diagram', 2)
            if er_diag.get('entities'):
                for entity in self._ensure_list(er_diag.get('entities', [])):
                    if isinstance(entity, dict):
                        doc.add_paragraph(f"Entity: {entity.get('name', '')}")
                        if entity.get('attributes'):
                            doc.add_paragraph(f"  Attributes: {', '.join(self._ensure_list(entity.get('attributes', [])))}")
        
        # Technology Stack
        tech_stack = self._ensure_dict(project_data.get('technology_stack', {}))
        if tech_stack:
            doc.add_heading('4.5 Technology Stack', 2)
            
            if isinstance(tech_stack.get('frontend'), dict):
                frontend = tech_stack.get('frontend', {})
                doc.add_paragraph(f"Frontend: {', '.join(self._ensure_list(frontend.get('technologies', [])))}")
            elif tech_stack.get('frontend'):
                doc.add_paragraph(f"Frontend: {tech_stack.get('frontend')}")
            
            if isinstance(tech_stack.get('backend'), dict):
                backend = tech_stack.get('backend', {})
                doc.add_paragraph(f"Backend: {', '.join(self._ensure_list(backend.get('technologies', [])))}")
            elif tech_stack.get('backend'):
                doc.add_paragraph(f"Backend: {tech_stack.get('backend')}")
            
            if isinstance(tech_stack.get('database'), dict):
                database = tech_stack.get('database', {})
                doc.add_paragraph(f"Database: {database.get('type', 'N/A')}")
            elif tech_stack.get('database'):
                doc.add_paragraph(f"Database: {tech_stack.get('database')}")
        
        # System Requirements
        sys_req = self._ensure_dict(project_data.get('system_requirements', {}))
        if sys_req:
            doc.add_heading('4.6 System Requirements', 2)
            
            hw_req = self._ensure_list(sys_req.get('hardware', []))
            if hw_req:
                doc.add_heading('Hardware Requirements:', 3)
                for hw in hw_req:
                    if isinstance(hw, dict):
                        doc.add_paragraph(f"• {hw.get('component', '')}: {hw.get('specification', '')}")
                    else:
                        doc.add_paragraph(f"• {hw}", style='List Bullet')
            
            sw_req = self._ensure_list(sys_req.get('software', []))
            if sw_req:
                doc.add_heading('Software Requirements:', 3)
                for sw in sw_req:
                    if isinstance(sw, dict):
                        doc.add_paragraph(f"• {sw.get('name', '')}: {sw.get('version', '')} - {sw.get('purpose', '')}")
                    else:
                        doc.add_paragraph(f"• {sw}", style='List Bullet')
        
        # ============ CHAPTER 5: IMPLEMENTATION ============
        doc.add_page_break()
        doc.add_heading('CHAPTER 5: IMPLEMENTATION', 1)
        
        # Modules
        modules = self._ensure_list(project_data.get('modules', []))
        if modules:
            doc.add_heading('5.1 Project Modules', 2)
            for idx, module in enumerate(modules, 1):
                if isinstance(module, str):
                    doc.add_heading(f"5.1.{idx} {module}", 3)
                elif isinstance(module, dict):
                    doc.add_heading(f"5.1.{idx} {module.get('name', 'Module')}", 3)
                    doc.add_paragraph(module.get('description', ''))
                    
                    if module.get('functionality'):
                        doc.add_paragraph("Functionality:")
                        for func in self._ensure_list(module.get('functionality', [])):
                            doc.add_paragraph(f"  • {func}")
                    
                    if module.get('input'):
                        doc.add_paragraph(f"Input: {module.get('input')}")
                    if module.get('output'):
                        doc.add_paragraph(f"Output: {module.get('output')}")
                    if module.get('duration_weeks'):
                        doc.add_paragraph(f"Duration: {module.get('duration_weeks')} weeks")
        
        # Database Design
        db_design = self._ensure_dict(project_data.get('database_design', {}))
        if db_design:
            doc.add_heading('5.2 Database Design', 2)
            
            if db_design.get('schema_description'):
                doc.add_paragraph(db_design.get('schema_description'))
            
            if db_design.get('normalization'):
                doc.add_paragraph(f"Normalization: {db_design.get('normalization')}")
            
            tables = self._ensure_list(db_design.get('tables', []))
            if tables:
                doc.add_heading('Database Tables:', 3)
                for table in tables:
                    if isinstance(table, dict):
                        doc.add_paragraph(f"Table: {table.get('name', 'table_name')}")
                        if table.get('description'):
                            doc.add_paragraph(f"  Purpose: {table.get('description')}")
                        
                        columns = self._ensure_list(table.get('columns', []))
                        if columns:
                            # Create columns table
                            col_headers = ["Column Name", "Data Type", "Constraints"]
                            col_rows = []
                            for col in columns:
                                if isinstance(col, dict):
                                    col_rows.append([
                                        col.get('name', ''),
                                        col.get('type', ''),
                                        col.get('constraints', '')
                                    ])
                                elif isinstance(col, str):
                                    col_rows.append([col, '', ''])
                            
                            if col_rows:
                                self._add_table_from_data(doc, col_headers, col_rows)
                                doc.add_paragraph()
        
        # Implementation Details
        impl = self._ensure_dict(project_data.get('implementation', {}))
        if impl and isinstance(impl, dict):
            doc.add_heading('5.3 Implementation Details', 2)
            
            if impl.get('development_environment'):
                doc.add_paragraph(f"Development Environment: {impl.get('development_environment')}")
            if impl.get('coding_standards'):
                doc.add_paragraph(f"Coding Standards: {impl.get('coding_standards')}")
            if impl.get('folder_structure'):
                doc.add_paragraph(f"Folder Structure: {impl.get('folder_structure')}")
        
        # ============ CHAPTER 6: TESTING ============
        doc.add_page_break()
        doc.add_heading('CHAPTER 6: TESTING', 1)
        
        testing = self._ensure_dict(project_data.get('testing', {}))
        if isinstance(project_data.get('testing'), str):
            doc.add_paragraph(project_data.get('testing'))
        else:
            if testing.get('testing_methodology'):
                doc.add_heading('6.1 Testing Methodology', 2)
                doc.add_paragraph(testing.get('testing_methodology'))
            
            if testing.get('tools_used'):
                doc.add_paragraph(f"Tools Used: {', '.join(self._ensure_list(testing.get('tools_used', [])))}")
            
            test_cases = self._ensure_list(testing.get('test_cases', []))
            if test_cases:
                doc.add_heading('6.2 Test Cases', 2)
                
                tc_headers = ["TC ID", "Module", "Test Scenario", "Expected Result", "Status"]
                tc_rows = []
                
                for tc in test_cases:
                    if isinstance(tc, dict):
                        tc_rows.append([
                            tc.get('tc_id', ''),
                            tc.get('module', ''),
                            tc.get('test_scenario', ''),
                            tc.get('expected_result', ''),
                            tc.get('status', 'Pass')
                        ])
                
                if tc_rows:
                    self._add_table_from_data(doc, tc_headers, tc_rows)
            
            test_summary = self._ensure_dict(testing.get('test_summary', {}))
            if test_summary:
                doc.add_heading('6.3 Test Summary', 2)
                doc.add_paragraph(f"Total Test Cases: {test_summary.get('total_cases', 'N/A')}")
                doc.add_paragraph(f"Passed: {test_summary.get('passed', 'N/A')}")
                doc.add_paragraph(f"Failed: {test_summary.get('failed', 'N/A')}")
                doc.add_paragraph(f"Pass Percentage: {test_summary.get('pass_percentage', 'N/A')}")
        
        # ============ CHAPTER 7: RESULTS & SCREENSHOTS ============
        doc.add_page_break()
        doc.add_heading('CHAPTER 7: RESULTS & SCREENSHOTS', 1)
        
        screenshots = self._ensure_list(project_data.get('screenshots', []))
        if screenshots:
            for idx, screen in enumerate(screenshots, 1):
                if isinstance(screen, dict):
                    doc.add_heading(f"7.{idx} {screen.get('screen_name', 'Screen')}", 2)
                    doc.add_paragraph(screen.get('description', ''))
                    if screen.get('key_features'):
                        doc.add_paragraph("Key Features:")
                        for feature in self._ensure_list(screen.get('key_features', [])):
                            doc.add_paragraph(f"  • {feature}")
                    doc.add_paragraph("[Screenshot Placeholder]")
                    doc.add_paragraph()
        else:
            doc.add_paragraph("Screenshots to be added based on implementation.")
        
        # ============ CHAPTER 8: CONCLUSION & FUTURE SCOPE ============
        doc.add_page_break()
        doc.add_heading('CHAPTER 8: CONCLUSION & FUTURE SCOPE', 1)
        
        doc.add_heading('8.1 Conclusion', 2)
        conclusion = project_data.get('conclusion', '')
        doc.add_paragraph(conclusion if conclusion else "Project conclusion to be written.")
        
        doc.add_heading('8.2 Future Scope', 2)
        future_scope = project_data.get('future_scope', [])
        if isinstance(future_scope, str):
            doc.add_paragraph(future_scope)
        elif isinstance(future_scope, list):
            for item in future_scope:
                doc.add_paragraph(item, style='List Bullet')
        
        # Limitations
        scope_data = self._ensure_dict(project_data.get('scope', {}))
        limitations = self._ensure_list(scope_data.get('limitations', []))
        if limitations:
            doc.add_heading('8.3 Limitations', 2)
            for lim in limitations:
                doc.add_paragraph(lim, style='List Bullet')
        
        # ============ REFERENCES ============
        doc.add_page_break()
        doc.add_heading('REFERENCES', 1)
        
        references = self._ensure_list(project_data.get('references', []))
        if references:
            for idx, ref in enumerate(references, 1):
                if isinstance(ref, str):
                    doc.add_paragraph(f"[{idx}] {ref}")
                elif isinstance(ref, dict):
                    citation = ref.get('citation', '')
                    ref_type = ref.get('type', '')
                    url = ref.get('url', '')
                    doc.add_paragraph(f"[{idx}] {citation}")
                    if url:
                        doc.add_paragraph(f"    URL: {url}")
        else:
            doc.add_paragraph("[1] References to be added.")
        
        # ============ APPENDIX: CODE SAMPLES ============
        code_samples = self._ensure_list(project_data.get('code_snippets', [])) or self._ensure_list(project_data.get('code_samples', []))
        if code_samples:
            doc.add_page_break()
            doc.add_heading('APPENDIX A: SOURCE CODE', 1)
            
            for idx, snippet in enumerate(code_samples, 1):
                if isinstance(snippet, str):
                    doc.add_paragraph(snippet)
                elif isinstance(snippet, dict):
                    doc.add_heading(f"A.{idx} {snippet.get('filename', 'Code File')}", 2)
                    if snippet.get('purpose'):
                        doc.add_paragraph(f"Purpose: {snippet.get('purpose')}")
                    
                    code_para = doc.add_paragraph()
                    code_content = snippet.get('content', snippet.get('code', ''))
                    code_run = code_para.add_run(code_content)
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(9)
                    
                    if snippet.get('explanation'):
                        doc.add_paragraph()
                        exp_para = doc.add_paragraph()
                        exp_para.add_run("Explanation: ").bold = True
                        exp_para.add_run(snippet.get('explanation'))
        
        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream.read()


# Singleton instance
docx_generator = DOCXGenerator()
